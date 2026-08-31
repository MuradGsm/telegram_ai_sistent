import io

import pypdf
from docx import Document as DocxDocument

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
TXT_CONTENT_TYPE = "text/plain"

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100


def extract_text(file_bytes: bytes, content_type: str) -> str:
    if content_type == PDF_CONTENT_TYPE:
        return _extract_pdf(file_bytes)
    if content_type == DOCX_CONTENT_TYPE:
        return _extract_docx(file_bytes)
    if content_type == TXT_CONTENT_TYPE:
        return file_bytes.decode("utf-8", errors="ignore")

    raise ValueError(f"Unsupported content type for text extraction: {content_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(page_texts)


def _extract_docx(file_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Извлекаем параграфы
    for paragraph in document.paragraphs:
        if text := paragraph.text.strip():
            parts.append(text)

    # Извлекаем текст из таблиц (важно для прайсов и инструкций)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def split_into_chunks(
    text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    text = text.strip()
    if not text:
        return []

    # Защита от бесконечного цикла при ошибочных оверлапах
    step = max(1, chunk_size - overlap)
    chunks: list[str] = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        
        # Стараемся не рвать слова: передвигаем 'end' до ближайшего пробела, если мы не в самом конце
        if end < text_length:
            boundary = text.rfind(" ", start, end)
            if boundary > start:
                end = boundary

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start += step

    return chunks